from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, italic=False, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.line_spacing = Pt(18)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p

def chapter(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = Pt(18)
    r = p.add_run(text.upper())
    set_font(r, size=14, bold=True)

def section(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    r = p.add_run(text)
    set_font(r, size=12, bold=True)

def subsection(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    r = p.add_run(text)
    set_font(r, size=12, bold=True, italic=True)

def placeholder(num, label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEEEEE')
    p._p.get_or_add_pPr().append(shd)
    r = p.add_run(f'[Figure {num}: {label} — insert diagram/screenshot here]')
    set_font(r, size=11, italic=True, color=(100,100,100))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(f'Figure {num}: {label}')
    set_font(rc, size=10, italic=True)

def tbl(headers, rows, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = cp.add_run(caption)
        set_font(r, size=11, bold=True)
        cp.paragraph_format.space_after = Pt(3)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=10, bold=True)
        p.paragraph_format.space_after = Pt(2)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, size=10)
            p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

def use_case_tbl(num, name, actor, precond, main_flow, alt_flow, postcond):
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = cp.add_run(f'Table {num}: Use Case Description – {name}')
    set_font(r, size=11, bold=True)
    cp.paragraph_format.space_after = Pt(3)
    rows_data = [
        ('Use Case Name', name),
        ('Actor(s)', actor),
        ('Precondition', precond),
        ('Main Flow', main_flow),
        ('Alternative Flow', alt_flow),
        ('Postcondition', postcond),
    ]
    t = doc.add_table(rows=len(rows_data)+1, cols=2)
    t.style = 'Table Grid'
    hcells = t.rows[0].cells
    for txt, cell in zip(['Field', 'Description'], hcells):
        r2 = cell.paragraphs[0].add_run(txt)
        set_font(r2, size=10, bold=True)
    for ri, (field, val) in enumerate(rows_data):
        cells = t.rows[ri+1].cells
        r3 = cells[0].paragraphs[0].add_run(field)
        set_font(r3, size=10, bold=True)
        r4 = cells[1].paragraphs[0].add_run(val)
        set_font(r4, size=10)
    doc.add_paragraph()

# ===================== TITLE PAGE =====================
para('Adventist University of Central Africa', WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True)
para('Faculty of Information Technology', WD_ALIGN_PARAGRAPH.CENTER, 12, bold=True)
para('Department of Software Engineering', WD_ALIGN_PARAGRAPH.CENTER, 12, bold=True)
doc.add_paragraph()
doc.add_paragraph()
para('DESIGN AND IMPLEMENTATION OF AN AI-POWERED BUYER-SELLER MATCHMAKING SYSTEM FOR PET TRADE AND SERVICES', WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True)
doc.add_paragraph()
para('Case Study: Hemmy Kennel, Lagos, Nigeria', WD_ALIGN_PARAGRAPH.CENTER, 12, bold=True)
doc.add_paragraph()
doc.add_paragraph()
para('A Final Year Project Presented in partial fulfillment of the requirements for the degree of', WD_ALIGN_PARAGRAPH.CENTER, 12)
para('BACHELOR OF SCIENCE IN INFORMATION TECHNOLOGY', WD_ALIGN_PARAGRAPH.CENTER, 12, bold=True)
para('Major in Software Engineering', WD_ALIGN_PARAGRAPH.CENTER, 12)
doc.add_paragraph()
doc.add_paragraph()
para('By', WD_ALIGN_PARAGRAPH.CENTER, 12)
para('FAMILONI Emmanuel Eniola', WD_ALIGN_PARAGRAPH.CENTER, 12, bold=True)
para('Student ID: 25951', WD_ALIGN_PARAGRAPH.CENTER, 12)
doc.add_paragraph()
para('Supervisor: Ishimwe Prince', WD_ALIGN_PARAGRAPH.CENTER, 12)
para('Tel: +250792525543', WD_ALIGN_PARAGRAPH.CENTER, 12)
doc.add_paragraph()
para('June 2026', WD_ALIGN_PARAGRAPH.CENTER, 12, bold=True)
doc.add_page_break()

# ===================== ABSTRACT =====================
chapter('Abstract')
para('The pet trade and breeding industry is growing rapidly on a global scale, with Nigeria emerging as a significant market driven by increasing urbanisation and rising demand for quality companion animals. Hemmy Kennel, a pet breeding and trading business based in Lagos, Nigeria, currently relies entirely on manual processes for buyer-seller matching, communicating through WhatsApp, Instagram, and phone calls without any structured platform. This manual approach limits efficiency, personalization, and the ability to scale in response to growing demand. This project presents the design and implementation of PetMatchAI, an AI-powered buyer-seller matchmaking system that intelligently connects pet buyers and sellers based on personalized preferences and behavioral data.')
para('The system leverages the Groq LLM API (llama-3.3-70b-versatile) for AI-powered match scoring and recommendation generation, Next.js 16 App Router for the frontend, Supabase (PostgreSQL) for the backend database, authentication, real-time subscriptions, and file storage, and a comprehensive multi-channel notification system encompassing in-app notifications, email via Brevo SMTP, SMS via Termii for Nigerian phone numbers, and browser push notifications via the Web Push API with VAPID keys.')
para('PetMatchAI comprises thirteen fully integrated functional modules: (1) User Registration and Authentication with Google OAuth, email OTP verification, login attempt lockout, and 30-minute idle session timeout; (2) Role-based Dashboards for buyers, sellers, and administrators; (3) Pet Listing Management with multi-photo upload, video support, and health documentation; (4) Intelligent Search and Discovery with breed autocomplete, map-based location search, and pet comparison; (5) AI Recommendation Engine with transparent match percentage scoring and contextual recommendations; (6) Buyer-Seller Communication with real-time messaging, typing indicators, read receipts, and Jitsi Meet video calling; (7) Matchmaking Engine with accept/decline, bidirectional matching, and history; (8) Analytics and Decision Support with KPI cards, charts, and AI-generated executive summaries; (9) Review and Rating with multi-criteria scoring; (10) Multi-channel Notifications; (11) Extended Reporting with Excel and PDF export; (12) Platform Administration with KYC verification and dispute resolution; and (13) Security and Audit with comprehensive audit logging and role-based access control.')
para('The system was developed using an incremental methodology, with data collected through structured interviews with Hemmy Kennel staff. Testing across all thirteen modules confirmed that the system meets all functional and non-functional requirements. AI match scoring returns results in under one second, real-time messages are delivered in an average of 0.3 seconds, and all security controls — including JWT authentication, Row Level Security, and identity verification — function as specified. The implemented system represents a complete digital transformation of Hemmy Kennel\'s buyer-seller engagement model.')
para('Keywords: Artificial Intelligence, Matchmaking System, Pet Trade, Recommendation Engine, Next.js, Supabase, Groq LLM, Buyer-Seller Platform.', italic=True)
doc.add_page_break()

# ===================== DECLARATION =====================
chapter('Declaration')
para('I, FAMILONI Emmanuel Eniola, Student ID Number 25951, am a registered student at the Adventist University of Central Africa in the Faculty of Information Technology, Department of Software Engineering. I hereby declare that this final year project report is entirely the product of my own original work and effort, to the best of my knowledge and belief. All sources consulted and referenced in this work have been duly acknowledged. This project report has not been submitted, either partially or wholly, to any other university or institution of higher learning for any academic award or qualification.')
doc.add_paragraph()
para('Signature: ________________________________      Date: ____________________')
doc.add_paragraph()
para('FAMILONI Emmanuel Eniola')
para('Student ID: 25951')
doc.add_page_break()

# ===================== APPROVAL =====================
chapter('Approval')
para('I, Ishimwe Prince, hereby certify that this final year project entitled "Design and Implementation of an AI-Powered Buyer-Seller Matchmaking System for Pet Trade and Services" has been carried out by FAMILONI Emmanuel Eniola (Student ID: 25951) under my supervision. I confirm that this work has been completed to a satisfactory standard and is submitted with my approval as the academic supervisor.')
doc.add_paragraph()
para('Supervisor Name: Ishimwe Prince')
doc.add_paragraph()
para('Signature: ________________________________      Date: ____________________')
doc.add_paragraph()
para('Faculty of Information Technology')
para('Department of Software Engineering')
para('Adventist University of Central Africa (AUCA)')
doc.add_page_break()

# ===================== DEDICATION =====================
chapter('Dedication')
para('This final year project is dedicated, first and foremost, to Almighty God, whose grace, wisdom, and provision have sustained me throughout every stage of this academic journey.')
para('To my parents, for their unconditional love, unwavering support, prayers, and sacrifices that made my education possible. Your belief in me has been the foundation of everything I have accomplished.')
para('To my supervisor, Ishimwe Prince, for his patient guidance, constructive feedback, and consistent encouragement throughout this project. Your mentorship has shaped this work in immeasurable ways.')
para('To my classmates and colleagues at AUCA, for the shared learning, encouragement, and companionship that made this academic journey both challenging and deeply rewarding.')
doc.add_page_break()

# ===================== ACKNOWLEDGMENTS =====================
chapter('Acknowledgments')
para('First and foremost, I give all glory and honour to Almighty God for His endless grace and wisdom that guided me through every challenge encountered during this research. His strength sustained me in difficult moments and His provision ensured that I had all the resources necessary to complete this work.')
para('I extend my deepest and most sincere gratitude to my supervisor, Ishimwe Prince, for his invaluable guidance, patience, and academic mentorship throughout the course of this project. His constructive feedback, expert insights, and consistent availability for consultation were instrumental in shaping the direction and quality of this research. I am truly grateful for the time and dedication he invested in my academic development.')
para('I am profoundly grateful to the academic staff of the Faculty of Information Technology and the Department of Software Engineering at the Adventist University of Central Africa (AUCA). The knowledge, skills, and values imparted through your teaching have equipped me with the technical competencies required to undertake and complete a project of this complexity.')
para('My heartfelt appreciation goes to the management and staff of Hemmy Kennel, Lagos, Nigeria, for their cooperation and openness during the data collection phase. Their willingness to share information about their operational challenges provided the real-world grounding that this research required and ensured the system developed is practically relevant.')
para('To my parents, thank you for your love, prayers, financial support, and belief in my abilities. Your sacrifices have made my education and this achievement possible. I am eternally grateful.')
para('To my classmates and friends at AUCA, thank you for your encouragement, collaborative spirit, and the many study sessions that sharpened my understanding. Your support throughout this journey has been invaluable.')
doc.add_paragraph()
para('FAMILONI Emmanuel Eniola')
doc.add_page_break()

# ===================== TABLE OF CONTENTS =====================
chapter('Table of Contents')
toc_entries = [
    ('ABSTRACT', 'ii'),
    ('DECLARATION', 'iii'),
    ('APPROVAL', 'iv'),
    ('DEDICATION', 'v'),
    ('ACKNOWLEDGMENTS', 'vi'),
    ('TABLE OF CONTENTS', 'vii'),
    ('LIST OF FIGURES', 'ix'),
    ('LIST OF TABLES', 'x'),
    ('LIST OF ABBREVIATIONS', 'xi'),
    ('CHAPTER 1: GENERAL INTRODUCTION', '1'),
    ('    Introduction', '1'),
    ('    Background of the Study', '2'),
    ('    Statement of the Problem', '4'),
    ('    Choice and Motivation of the Study', '6'),
    ('    Objectives of the Study', '7'),
    ('    Scope of the Study', '8'),
    ('    Methodology and Techniques Used', '9'),
    ('    Expected Results', '11'),
    ('    Organization of the Work', '12'),
    ('CHAPTER 2: ANALYSIS OF THE EXISTING SYSTEM', '13'),
    ('    Introduction', '13'),
    ('    Description of the Current System Environment', '13'),
    ('    Description of the Current System', '15'),
    ('    Analysis of the Current System', '17'),
    ('    Modeling of the Current System', '19'),
    ('    Problems of the Current System', '20'),
    ('    Proposed Solutions', '22'),
    ('    System Requirements', '23'),
    ('CHAPTER 3: REQUIREMENTS ANALYSIS AND DESIGN', '26'),
    ('    Introduction', '26'),
    ('    Analysis and Design Methodology', '26'),
    ('    Use Case Diagram', '27'),
    ('    Use Case Descriptions', '28'),
    ('    Class Diagram', '36'),
    ('    Sequence Diagrams', '37'),
    ('    Activity Diagram', '41'),
    ('    Database Diagram (ERD)', '42'),
    ('    Data Dictionary', '44'),
    ('    System Architectural Design', '52'),
    ('CHAPTER 4: IMPLEMENTATION OF THE NEW SYSTEM', '55'),
    ('    Introduction', '55'),
    ('    Technologies and Tools Used', '55'),
    ('    Screenshots of the Implemented System', '60'),
    ('    Software Testing', '70'),
    ('    Hardware and Software Requirements', '73'),
    ('CHAPTER 5: CONCLUSION AND RECOMMENDATIONS', '75'),
    ('    Conclusion', '75'),
    ('    Recommendations', '77'),
    ('REFERENCES', '79'),
    ('APPENDICES', '82'),
    ('    Appendix A: Curriculum Vitae', '82'),
    ('    Appendix B: Data Collection Authorization Letter', '84'),
]
for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    is_chapter = entry.startswith('CHAPTER') or entry in ('ABSTRACT','DECLARATION','APPROVAL','DEDICATION','ACKNOWLEDGMENTS','TABLE OF CONTENTS','LIST OF FIGURES','LIST OF TABLES','LIST OF ABBREVIATIONS','REFERENCES','APPENDICES')
    r = p.add_run(entry)
    set_font(r, size=11, bold=is_chapter)
    # dots and page number
    tab = p.add_run(f'\t{page}')
    set_font(tab, size=11, bold=is_chapter)
doc.add_page_break()

# ===================== LIST OF FIGURES =====================
chapter('List of Figures')
figures = [
    ('Figure 1', 'Modeling of the Current Manual Process at Hemmy Kennel'),
    ('Figure 2', 'Entity Relationship Diagram (ERD)'),
    ('Figure 3', 'Use Case Diagram'),
    ('Figure 4', 'Class Diagram'),
    ('Figure 5', 'Sequence Diagram: User Registration'),
    ('Figure 6', 'Sequence Diagram: AI Matching Flow'),
    ('Figure 7', 'Sequence Diagram: Buyer-Seller Messaging'),
    ('Figure 8', 'Activity Diagram: Main User Flow'),
    ('Figure 9', 'System Architecture Diagram'),
    ('Figure 10', 'Screenshot: Login Page'),
    ('Figure 11', 'Screenshot: Buyer Dashboard'),
    ('Figure 12', 'Screenshot: Pet Listings Page'),
    ('Figure 13', 'Screenshot: AI Recommendations Page'),
    ('Figure 14', 'Screenshot: Matchmaking Page'),
    ('Figure 15', 'Screenshot: Messages Page'),
    ('Figure 16', 'Screenshot: Analytics Dashboard'),
    ('Figure 17', 'Screenshot: Admin Panel'),
    ('Figure 18', 'Screenshot: Profile and Identity Verification'),
    ('Figure 19', 'Screenshot: Offers and Dispute Filing'),
]
for fig, desc in figures:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(f'{fig}: {desc}')
    set_font(r, size=11)
doc.add_page_break()

# ===================== LIST OF TABLES =====================
chapter('List of Tables')
tables_list = [
    ('Table 1', 'Use Case Description: Register User Account'),
    ('Table 2', 'Use Case Description: Login to System'),
    ('Table 3', 'Use Case Description: Create Pet Listing'),
    ('Table 4', 'Use Case Description: Get AI Recommendations'),
    ('Table 5', 'Use Case Description: Send Message'),
    ('Table 6', 'Use Case Description: Submit Price Offer'),
    ('Table 7', 'Use Case Description: File Dispute'),
    ('Table 8', 'Data Dictionary: profiles Table'),
    ('Table 9', 'Data Dictionary: pets Table'),
    ('Table 10', 'Data Dictionary: buyer_preferences Table'),
    ('Table 11', 'Data Dictionary: ai_matches Table'),
    ('Table 12', 'Data Dictionary: messages Table'),
    ('Table 13', 'Data Dictionary: offers Table'),
    ('Table 14', 'Data Dictionary: reviews Table'),
    ('Table 15', 'Data Dictionary: notifications Table'),
    ('Table 16', 'Data Dictionary: disputes Table'),
    ('Table 17', 'Data Dictionary: user_verification_requests Table'),
    ('Table 18', 'Functional Requirements'),
    ('Table 19', 'Non-Functional Requirements'),
    ('Table 20', 'Hardware and Software Requirements'),
    ('Table 21', 'Software Testing Results'),
]
for tbl_ref, desc in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(f'{tbl_ref}: {desc}')
    set_font(r, size=11)
doc.add_page_break()

# ===================== LIST OF ABBREVIATIONS =====================
chapter('List of Abbreviations')
abbrevs = [
    ('AI', 'Artificial Intelligence'),
    ('API', 'Application Programming Interface'),
    ('AUCA', 'Adventist University of Central Africa'),
    ('CSS', 'Cascading Style Sheets'),
    ('FK', 'Foreign Key'),
    ('HTTP', 'Hypertext Transfer Protocol'),
    ('HTTPS', 'Hypertext Transfer Protocol Secure'),
    ('ID', 'Identifier'),
    ('IT', 'Information Technology'),
    ('JSON', 'JavaScript Object Notation'),
    ('JWT', 'JSON Web Token'),
    ('KYC', 'Know Your Customer'),
    ('LLM', 'Large Language Model'),
    ('MFA', 'Multi-Factor Authentication'),
    ('OTP', 'One-Time Password'),
    ('PK', 'Primary Key'),
    ('RLS', 'Row Level Security'),
    ('SMS', 'Short Message Service'),
    ('SQL', 'Structured Query Language'),
    ('UI', 'User Interface'),
    ('UML', 'Unified Modeling Language'),
    ('UX', 'User Experience'),
    ('UUID', 'Universally Unique Identifier'),
    ('VAPID', 'Voluntary Application Server Identification'),
    ('WebRTC', 'Web Real-Time Communication'),
]
for abbr, meaning in abbrevs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(16)
    r1 = p.add_run(f'{abbr}')
    set_font(r1, size=11, bold=True)
    r2 = p.add_run(f' - {meaning}')
    set_font(r2, size=11)
doc.add_page_break()

# ===================== CHAPTER 1 =====================
chapter('Chapter 1')
chapter('General Introduction')

section('Introduction')
para('The global pet industry has undergone remarkable transformation in recent decades, evolving from informal local markets into a sophisticated and technology-enabled ecosystem encompassing pet trading, breeding, healthcare, accessories, and services. According to the American Pet Products Association, global pet industry expenditure surpassed $150 billion in 2023 and continues to grow at an accelerating pace driven by urbanisation, changing household structures, and deepening human-animal bonds. In Nigeria, this growth trajectory is equally pronounced, with increasing numbers of urban households acquiring dogs, cats, and other companion animals, creating significant demand for structured, trustworthy, and efficient platforms that can connect buyers with quality sellers.')
para('Artificial intelligence has emerged as a transformative force in digital marketplace platforms, enabling personalised product discovery, intelligent matchmaking, behavioural analytics, and automated customer engagement at scale. Leading e-commerce platforms such as Amazon and Netflix have demonstrated the commercial and user-experience value of AI-powered recommendation engines, inspiring their adoption across diverse sectors including real estate, healthcare, and financial services. In the context of pet trading, AI offers compelling opportunities to automate buyer-seller matching, personalise pet discovery, and generate business insights from accumulated preference and behavioural data.')
para('This project presents PetMatchAI, an AI-powered buyer-seller matchmaking system designed to intelligently connect pet buyers and sellers at Hemmy Kennel, a pet breeding and trading business based in Lagos, Nigeria. The system replaces the current manual engagement model — which relies on WhatsApp messaging, Instagram posts, and telephone calls — with a centralised, intelligent platform featuring thirteen integrated functional modules that cover the complete pet trade lifecycle from buyer registration through listing discovery, AI-powered matching, real-time communication, offer negotiation, transaction review, and business analytics.')

section('Background of the Study')
para('The global pet trade industry is characterised by fragmentation, geographic diversity, and high variability in product quality and seller credibility. Buyers seeking specific breeds, health certifications, and price points face significant challenges in discovering suitable options through traditional channels, which lack filtering, personalisation, and transparent quality indicators. The rise of social media has partially addressed discoverability but introduced new challenges around trust, information completeness, and communication efficiency.')
para('In Nigeria, the pet market is concentrated in major urban centres, particularly Lagos, Abuja, and Port Harcourt, where demand for purebred dogs and cats has grown significantly among middle and upper-income households. Despite this growth, most pet sellers in Nigeria continue to operate through informal channels. Hemmy Kennel, a Lagos-based pet breeding and trading business, exemplifies this operational model. The business maintains a portfolio of dogs and cats spanning multiple breeds, but relies on WhatsApp status updates, Instagram posts, and word-of-mouth referrals to reach potential buyers. Buyer inquiries arrive through WhatsApp messages and phone calls, are handled individually by staff members, and are not tracked in any centralised system.')
para('Artificial intelligence and machine learning technologies have demonstrated transformative potential in addressing the buyer-seller matching challenge in diverse markets. Collaborative filtering approaches, pioneered by Resnick and Varian (1997) and later refined through deep learning by Covington, Adams, and Sargin (2016) at YouTube, enable platforms to generate personalised recommendations from user preference and behavioural signals at scale. Content-based filtering approaches complement behavioural techniques by matching structured product attributes against buyer preference profiles. Hybrid approaches combining both methods have shown superior performance in niche product categories (Burke, 2002), making them particularly suited to the specialised domain of pet matching where breed-specific attributes, health credentials, and geographic proximity are critical decision factors.')
para('Large language models (LLMs) represent the current frontier of AI-powered recommendation, offering the ability to generate natural language explanations for recommendations, synthesise complex attribute relationships, and produce personalised narrative content. The Groq API, built on the llama-3.3-70b-versatile model, provides ultra-fast LLM inference exceeding 400 tokens per second, enabling real-time AI scoring of pet listings against buyer preferences without degrading the user experience. This technological capability forms the foundation of PetMatchAI\'s recommendation engine.')
para('Research on preference elicitation by Chen and Pu (2012) established that structured preference capture at registration significantly improves recommendation accuracy compared to purely behavioural inference. This finding informed PetMatchAI\'s multi-step registration flow for buyers, which captures species preferences, budget range, purpose, and health requirements during account creation, providing immediate personalisation from the first interaction.')

section('Statement of the Problem')
para('The buyer-seller matching and pet trade process at Hemmy Kennel is characterised by a high degree of manual effort, limited personalisation, and inconsistent buyer experience. The following specific problems have been identified through direct observation and interviews with Hemmy Kennel management and staff:')
para('First, all buyer inquiries arrive through informal channels — WhatsApp messages, Instagram direct messages, and telephone calls — without any structured intake process. There is no centralised system for capturing, categorising, or tracking buyer requirements, resulting in incomplete information capture and frequent miscommunication about breed specifications, price expectations, and health requirements.')
para('Second, pet listings are published manually through social media posts, without standardised formats, consistent health documentation, or reliable pricing information. Buyers must scroll through all available posts to find potentially suitable pets, with no filtering, sorting, or personalised ranking capability. This undifferentiated discovery process results in high rates of irrelevant inquiry and buyer frustration.')
para('Third, sellers handle each inquiry individually, spending significant time responding to questions that could be answered through a structured listing profile. This manual inquiry management reduces the time available for animal care and breeding operations, limits the number of inquiries that can be handled simultaneously, and creates bottlenecks during periods of high demand.')
para('Fourth, there is no systematic capture of buyer preference and behavioural data. Hemmy Kennel cannot identify which breeds attract the most interest, what price points generate the most inquiries, or how buyer preferences vary by geography or season. This absence of business intelligence prevents data-driven decisions about breeding priorities, pricing strategy, and inventory management.')
para('Fifth, there is no structured communication platform for buyer-seller interaction beyond informal messaging. Offer negotiation, health documentation sharing, and transaction management occur through disconnected channels without any record-keeping or status tracking. Dispute resolution, if needed, has no formal process.')
para('Sixth, there is no review or rating system to build trust between buyers and sellers. New buyers have no way to assess seller credibility beyond informal social media reputation, limiting confidence and constraining the market to buyers already familiar with Hemmy Kennel.')

section('Choice and Motivation of the Study')
para('The decision to design and implement PetMatchAI for Hemmy Kennel is motivated by academic, organisational, and personal considerations.')
para('Academically, this project contributes to the growing body of research on AI applications in niche e-commerce and matchmaking domains. The pet trade sector presents unique matching challenges — biological diversity, health credential requirements, emotional buyer decision-making, and geographic proximity constraints — that extend recommendation system research beyond the well-studied domains of film, music, and general retail. By applying and adapting state-of-the-art LLM-based recommendation to this domain, this project enriches AUCA\'s research portfolio and provides a practical case study for future students interested in AI-powered marketplace development.')
para('Organisationally, Hemmy Kennel has expressed strong interest in this project, recognising that it directly addresses the operational inefficiencies that limit its capacity to serve growing buyer demand. The system developed through this research will be deployed as a functional platform capable of supporting Hemmy Kennel\'s commercial operations, providing immediate practical value beyond its academic contribution.')
para('Personally, this project provides the opportunity to apply academic competencies in artificial intelligence, full-stack web development, database design, and system architecture to a real-world problem. The combination of AI engineering, user experience design, and business analytics involved in building PetMatchAI represents a comprehensive demonstration of Software Engineering skills acquired at AUCA.')

section('Objectives of the Study')
subsection('General Objective')
para('To design and implement an AI-powered buyer-seller matchmaking system that intelligently connects pet buyers and sellers based on user preferences and behavioural data, enhancing buyer experience, operational efficiency, and data-driven decision-making at Hemmy Kennel, Lagos, Nigeria.')
subsection('Specific Objectives')
objectives = [
    'To develop a structured relational database for pets, buyers, sellers, and interactions, supporting intelligent matching, search, and analytics operations.',
    'To design and implement an AI-based recommendation engine using the Groq LLM API (llama-3.3-70b-versatile) that generates personalised pet matches with transparent scoring and natural language explanations.',
    'To implement intelligent filtering and search functionalities enabling precise pet discovery by breed, species, age, price, gender, location, and health status, with map-based geographic search.',
    'To develop a real-time communication system supporting direct buyer-seller messaging with typing indicators, read receipts, message templates, video calling via Jitsi Meet, and offer negotiation.',
    'To build a data analytics module providing pricing analysis, listing conversion metrics, breed popularity rankings, average time-to-sell statistics, and AI-generated executive summaries.',
    'To develop a responsive web interface accessible on desktop and mobile browsers, providing intuitive access to all platform features for buyers, sellers, and administrators.',
    'To ensure system security through role-based access control, JWT authentication, Row Level Security, identity verification via government ID upload, audit logging, and login attempt lockout.',
    'To implement a post-transaction review and rating system supporting multi-criteria feedback, seller responses, and aggregate rating displays to build trust and accountability.',
]
for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(obj)
    set_font(r, size=12)

section('Scope of the Study')
para('This project encompasses the comprehensive design, development, testing, and documentation of PetMatchAI, an AI-powered buyer-seller matchmaking platform specifically designed for Hemmy Kennel\'s operational context. The system scope covers thirteen integrated functional modules addressing all aspects of the pet trade lifecycle from initial buyer registration and preference capture through listing discovery, AI-powered matching, communication, transaction management, analytics, and platform administration.')
para('The buyer-facing components include a registration and preference capture interface, AI-powered personalised pet recommendations with match scoring, advanced search and filtering tools, map-based location search, saved listings, pet comparison, and a structured communication module. The seller-facing components include comprehensive listing management with multi-media support, inquiry management, offer and negotiation tools, identity verification, and listing performance analytics. The platform administrative components include user management, content moderation, KYC verification approval, dispute resolution, audit log access, and broadcast announcements.')
para('The project explicitly excludes direct integration with external payment processing gateways, physical delivery logistics systems, and multi-organisation multi-tenant capabilities. Live video streaming for virtual pet viewing is treated as a future enhancement; video calling is implemented through the Jitsi Meet API rather than a custom streaming solution. The scope covers requirements analysis and system design, user interface implementation, full-stack web application development, comprehensive testing, and complete system documentation.')

section('Methodology and Techniques Used in the Study')
para('This project employs an Incremental Software Development Model, in which the system is built and validated in iterative cycles, with each cycle delivering a functional increment of the overall platform. This approach was selected because it allows early feedback on core components, enables adjustment of subsequent increments based on testing outcomes, and manages the complexity of a thirteen-module system through structured decomposition.')
subsection('Observation')
para('Direct observation of the day-to-day buyer inquiry and pet matching operations at Hemmy Kennel was conducted to understand how staff currently manage buyer communications, listing presentations, and price negotiations. Observation revealed that a single staff member typically handles multiple WhatsApp conversations simultaneously, with no system for tracking inquiry status or buyer preferences across sessions.')
subsection('Documentation Review')
para('Existing documentation at Hemmy Kennel, including social media post archives, WhatsApp conversation records, and informal price lists, was reviewed to understand current listing formats, pricing practices, and communication patterns. This review identified the specific data fields required for a structured pet listing schema and informed the design of the database.')
subsection('Interviews')
para('Structured interviews were conducted with Hemmy Kennel management and two staff members. Key questions and summarised responses included:')
interview_data = [
    ('Q1: How are buyer inquiries currently managed?', 'All inquiries arrive through WhatsApp and phone calls. There is no system — each staff member handles their own conversations and there is no shared record of who asked about which pet or at what price.'),
    ('Q2: What challenges do buyers face when trying to find a specific pet?', 'Buyers often cannot find the exact breed they want because they have to scroll through all our posts. Prices are not always visible, and buyers do not know the health status of pets without asking.'),
    ('Q3: What difficulties arise from managing inquiries manually?', 'We lose track of inquiries, miss potential buyers when we are busy with animals, and sometimes give inconsistent prices for the same pet to different buyers.'),
    ('Q4: What features would a digital platform need to be useful?', 'It should let buyers search by breed and price, show health certificates, allow direct messaging, and help us see which breeds are most popular so we can plan our breeding.'),
    ('Q5: How important is trust and verification to buyers?', 'Very important. New buyers want to know if we are legitimate and if the pet is healthy. A platform that shows our credentials and lets buyers leave reviews would help a lot.'),
]
for q, a in interview_data:
    p_q = doc.add_paragraph()
    p_q.paragraph_format.line_spacing = Pt(18)
    p_q.paragraph_format.space_after = Pt(2)
    r = p_q.add_run(q)
    set_font(r, size=12, bold=True)
    p_a = doc.add_paragraph()
    p_a.paragraph_format.line_spacing = Pt(18)
    p_a.paragraph_format.space_after = Pt(6)
    ra = p_a.add_run(a)
    set_font(ra, size=12, italic=True)

section('Expected Results')
para('Upon successful implementation and deployment of PetMatchAI, the following outcomes are anticipated:')
results = [
    'Improved buyer-seller match rates through AI-powered personalised recommendations that align pet listings with individual buyer preferences and behavioural signals.',
    'Reduced inquiry response time through structured listing profiles that answer common buyer questions automatically, reducing the volume of repetitive manual inquiries.',
    'Enhanced buyer experience through intelligent search, breed-specific filtering, map-based location discovery, and transparent AI match scoring with natural language explanations.',
    'Centralised communication through an integrated messaging platform with real-time delivery, reducing reliance on WhatsApp and enabling structured offer management.',
    'Data-driven business intelligence through analytics dashboards that reveal popular breeds, pricing trends, listing conversion rates, and geographic demand patterns.',
    'Improved platform trust through identity-verified seller profiles, multi-criteria buyer reviews, and a structured dispute resolution process.',
    'Operational efficiency gains through automation of notification delivery, match scoring, audit logging, and weekly analytics reporting.',
]
for res in results:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(res)
    set_font(r, size=12)

section('Organization of the Work')
org = [
    ('Chapter One – General Introduction', 'Provides the background, problem statement, objectives, scope, methodology, and expected results of the study.'),
    ('Chapter Two – Analysis of the Existing System', 'Examines the current manual pet trading process at Hemmy Kennel, identifies its weaknesses, and presents the proposed solutions and system requirements.'),
    ('Chapter Three – Requirements Analysis and Design', 'Presents the UML diagrams, data dictionary, use case descriptions, entity relationship diagram, and system architecture design of PetMatchAI.'),
    ('Chapter Four – Implementation of the New System', 'Describes the technologies and tools used, presents screenshots of the implemented system, and documents the software testing results.'),
    ('Chapter Five – Conclusion and Recommendations', 'Summarises the research outcomes and provides recommendations for future development and enhancement of PetMatchAI.'),
]
for title, desc in org:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + ': ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)
doc.add_page_break()

# ===================== CHAPTER 2 =====================
chapter('Chapter 2')
chapter('Analysis of the Existing System')

section('Introduction')
para('The purpose of designing and implementing PetMatchAI is to replace and enhance the existing manual processes currently in use at Hemmy Kennel. To build an efficient, consistent, and intelligent system, it is essential to first analyse how the current model operates, understand its structural limitations, and identify the specific areas that require technological intervention. A thorough examination of the existing system establishes the empirical foundation upon which the proposed solution is justified, ensuring that the new platform is designed to address real, observed operational deficiencies rather than assumed ones. Without this groundwork, any system developed would risk automating inefficiency rather than eliminating it.')
para('In software engineering practice, the analysis of an existing system is a critical prerequisite to the design of any replacement or improvement. It allows the developer to document how work is presently carried out, who the actors involved are, what information flows between them, and where delays, errors, and bottlenecks occur. By studying the current model in detail, the developer is able to distinguish between aspects of the present process that are valuable and worth preserving and those that are deficient and must be redesigned. This evidence-based approach reduces the risk of project failure and ensures that the resulting system is genuinely aligned with the needs of its users and the operational objectives of the business.')
para('Accordingly, this chapter examines the current methods used by Hemmy Kennel for pet listing, buyer inquiry management, and buyer-seller matching. It begins by describing the operational environment in which the business functions, including its historical background, mission, and vision, before presenting a detailed account of how the current manual system operates across its three core areas of listing publication, inquiry handling, and price negotiation. The chapter then critically analyses the strengths and, more significantly, the weaknesses of this manual approach, drawing on direct observation and on the interviews conducted with Hemmy Kennel management and staff during the data collection phase.')
para('Building on this analysis, the chapter models the existing workflow as a sequence of operational steps in order to make its structure and limitations explicit. It subsequently consolidates the recurring difficulties of the current system into a set of clearly defined problems — covering performance, scalability, data loss, inconsistency, the absence of business intelligence, and the deficit of trust — that collectively justify the need for an AI-powered solution. Finally, the chapter presents the proposed solutions offered by PetMatchAI and specifies the functional and non-functional system requirements that the new platform must satisfy. Together, these sections provide a comprehensive understanding of the problem domain and form the bridge between the general introduction presented in Chapter One and the requirements analysis and design detailed in Chapter Three.')

section('Description of the Current System Environment')
subsection('Historical Background')
para('Hemmy Kennel is a pet breeding and trading business established in Lagos, Nigeria. The business specialises in the breeding, raising, and trading of dogs and cats, serving individual buyers, pet enthusiasts, and families across Lagos and other Nigerian cities. Since its establishment, Hemmy Kennel has built a reputation for quality animals and reliable service, primarily through word-of-mouth referrals and social media presence on platforms including Instagram and WhatsApp. As demand for quality pets has grown — particularly among Lagos\'s expanding urban middle class — the limitations of Hemmy Kennel\'s manual operational approach have become increasingly apparent, creating the operational context and justification for this project.')
subsection('Mission')
para('To provide high-quality, healthy, and well-documented pets to buyers across Nigeria through honest, personalised, and reliable service that builds lasting relationships between sellers and pet owners.')
subsection('Vision')
para('To become Nigeria\'s most trusted pet breeding and trading brand, recognised for exceptional animal quality, buyer satisfaction, and technology-enabled customer experience.')

section('Description of the Current System')
para('The current system for managing pet listings and buyer-seller matching at Hemmy Kennel is entirely manual, social media-driven, and heavily reliant on individual staff effort and informal communication. The system covers three core operational areas: listing publication, buyer inquiry management, and price negotiation.')
subsection('Listing Publication Process')
para('Pet listings are created and published manually through Hemmy Kennel\'s social media accounts. Staff members photograph available pets, compose informal captions describing breed, age, and price, and post these on Instagram and WhatsApp Status. There is no standardised format for listings — some posts include health certification details while others do not; some show prices while others invite buyers to enquire. Videos of pets are occasionally posted but are not systematically produced for all available animals. Listings are not indexed or searchable beyond the native social media platform search functions.')
subsection('Buyer Inquiry Management Process')
para('Buyers who discover available pets through social media contact Hemmy Kennel directly through WhatsApp messages, Instagram direct messages, or phone calls. Each inquiry is handled individually by the staff member who receives it, without any centralised record of who has enquired about which pet. Staff members maintain personal WhatsApp chats with potential buyers, respond to questions about breed details, health history, and availability, and share additional photos or videos on request. There is no system for tracking inquiry status, following up with buyers who expressed interest but did not purchase, or analysing patterns in buyer requests over time.')
subsection('Price Negotiation and Transaction Process')
para('Price negotiation occurs informally through WhatsApp conversations or telephone calls. There is no structured offer management system — buyers make informal price proposals, and sellers respond with counter-offers through messaging. If a price is agreed, the buyer typically visits Hemmy Kennel in person to inspect the pet and complete the transaction through cash or bank transfer. There is no digital record of completed transactions, no post-transaction review process, and no dispute resolution mechanism beyond direct communication between buyer and seller.')

section('Analysis of the Current System')
para('The current manual system presents several operational weaknesses that limit Hemmy Kennel\'s effectiveness, scalability, and competitiveness:')
weaknesses = [
    ('No centralised listing platform', 'Listings exist across multiple social media accounts with no unified search or filtering capability. Buyers must manually browse all posts to find potentially suitable pets, resulting in time-consuming discovery and high rates of irrelevant inquiry.'),
    ('No personalised matching', 'The system has no mechanism for matching specific buyer requirements — breed preferences, budget constraints, health requirements, geographic location — against available listings. All buyers receive the same undifferentiated listing feed.'),
    ('Manual and unscalable inquiry handling', 'Each buyer inquiry requires individual staff attention. During peak demand periods, staff cannot effectively manage multiple simultaneous conversations, leading to delayed responses, missed inquiries, and inconsistent buyer experiences.'),
    ('No structured health or registration documentation', 'Health certification and registration information is not consistently included in listings. Buyers must explicitly request this information, creating friction in the discovery process and reducing buyer confidence.'),
    ('No buyer behavioural data', 'The system captures no data on buyer preferences, browsing behaviour, or inquiry patterns. Hemmy Kennel cannot identify demand trends, popular breeds, or optimal pricing strategies based on accumulated interaction data.'),
    ('No offer management or transaction tracking', 'Offer negotiation occurs through unstructured messaging with no status tracking, history recording, or formalised agreement mechanism. This creates risks of misunderstanding and provides no documentation for dispute resolution.'),
    ('No review or trust mechanism', 'New buyers have no reliable means of assessing seller credibility beyond informal social media reputation. The absence of a structured review system limits buyer confidence and constrains the market.'),
]
for title, desc in weaknesses:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + ': ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

section('Modeling of the Current System')
placeholder(1, 'Modeling of the Current Manual Process at Hemmy Kennel')
para('The current process at Hemmy Kennel can be modelled through the following sequential steps:')
steps = [
    ('Step 1: Listing Publication', 'Staff photograph available pets and post informal listings on Instagram and WhatsApp Status. No standardised format is used.'),
    ('Step 2: Buyer Discovery', 'Potential buyers browse social media posts or learn of available pets through word-of-mouth referrals. No search or filter capability exists.'),
    ('Step 3: Informal Inquiry', 'Interested buyers send WhatsApp messages or Instagram DMs requesting more information about specific pets.'),
    ('Step 4: Manual Response', 'Staff members respond individually to each inquiry, sharing additional photos, health details, and price information through chat messages.'),
    ('Step 5: Price Negotiation', 'Buyer and seller negotiate price through informal chat messages or telephone calls. No offer management system exists.'),
    ('Step 6: Agreement and Payment', 'If price is agreed, buyer arranges a physical visit to inspect the pet and completes payment by cash or bank transfer.'),
    ('Step 7: Pet Collection', 'Buyer collects the pet in person. No digital transaction record or review process follows.'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + ': ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

section('Problems of the Current System')
problems = [
    ('Poor Performance and Throughput', 'The manual system cannot handle multiple simultaneous buyer inquiries efficiently. Each inquiry requires dedicated staff attention, creating bottlenecks during peak demand and resulting in delayed responses that may drive potential buyers to competitors.'),
    ('Insufficient Scalability', 'As Hemmy Kennel\'s buyer base grows, the manual approach does not scale. Increasing inquiry volume directly increases staff workload without any automation buffer, limiting the business\'s growth capacity.'),
    ('Data Loss and No Audit Trail', 'Buyer preferences, inquiry history, negotiation records, and transaction details exist only in individual staff members\' personal WhatsApp conversations. This data is lost when staff change phones, leave the business, or when conversations are deleted.'),
    ('Inconsistent Buyer Experience', 'Different staff members provide different levels of detail, response speed, and pricing information. Buyers may receive inconsistent experiences depending on which staff member handles their inquiry.'),
    ('No Business Intelligence', 'Without structured data capture, Hemmy Kennel cannot identify popular breeds, seasonal demand patterns, optimal price points, or buyer geographic distribution — limiting strategic planning.'),
    ('Trust Deficit', 'The absence of verified seller credentials, standardised health documentation, and a buyer review system reduces buyer confidence, particularly for first-time customers without prior social media familiarity with Hemmy Kennel.'),
]
for title, desc in problems:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + ': ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

section('Proposed Solutions')
para('PetMatchAI is proposed as a comprehensive digital replacement for Hemmy Kennel\'s manual process. The system addresses each identified problem through thirteen integrated modules:')
solutions = [
    'A centralised pet listing management module with standardised formats, health documentation, and multi-media support replaces informal social media posts.',
    'An AI-powered recommendation engine using the Groq LLM generates personalised matches for each buyer based on their captured preferences, eliminating undifferentiated listing browsing.',
    'A real-time messaging module with structured threads, offer management, and notification delivery replaces WhatsApp-based inquiry handling.',
    'A relational PostgreSQL database with Row Level Security captures all buyer preferences, listing views, messages, offers, and transactions in a structured, searchable, and auditable format.',
    'An analytics dashboard provides breed popularity rankings, pricing trends, conversion rates, and geographic demand data to support business intelligence.',
    'An identity verification module and buyer review system build platform trust through verified seller credentials and transparent feedback.',
]
for sol in solutions:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(sol)
    set_font(r, size=12)

section('System Requirements')
subsection('Functional Requirements')
tbl(
    ['#', 'Requirement', 'Description'],
    [
        ('FR-01', 'User Registration', 'System shall support buyer, seller, and administrator registration with role selection, email OTP verification, and preference capture.'),
        ('FR-02', 'Authentication', 'System shall support email/password login, Google OAuth, login lockout after 5 failed attempts, and 30-minute idle session timeout.'),
        ('FR-03', 'Pet Listing Management', 'Sellers shall create, edit, and delete listings with pet details, health info, multiple photos, optional video, and draft saving.'),
        ('FR-04', 'Search and Discovery', 'Buyers shall search by breed, species, price, gender, location with autocomplete, map view, and pet comparison.'),
        ('FR-05', 'AI Recommendations', 'System shall generate personalised pet recommendations with match percentage scores and natural language explanations.'),
        ('FR-06', 'Messaging', 'System shall provide real-time buyer-seller messaging with typing indicators, read receipts, message templates, and video calling.'),
        ('FR-07', 'Offer Management', 'Buyers shall submit price offers; sellers shall accept, reject, or counter; both parties notified on status change.'),
        ('FR-08', 'Review and Rating', 'Buyers shall leave multi-criteria reviews after accepted offers; sellers shall respond; ratings aggregated on profiles.'),
        ('FR-09', 'Notifications', 'System shall deliver notifications via in-app, email, SMS, and push channels with user-configurable preferences.'),
        ('FR-10', 'Analytics', 'System shall provide KPI dashboards, breed pricing charts, conversion metrics, AI executive summaries, and data export.'),
        ('FR-11', 'Identity Verification', 'Sellers shall upload government ID for KYC; admin approves; verified badge displayed on listings and profiles.'),
        ('FR-12', 'Dispute Resolution', 'Users shall file disputes on accepted offers; admin resolves through admin panel with resolution notes.'),
        ('FR-13', 'Audit Logging', 'All user actions shall be logged to audit_logs table; admin shall filter and export logs.'),
    ],
    caption='Table 18: Functional Requirements'
)
subsection('Non-Functional Requirements')
tbl(
    ['#', 'Requirement', 'Specification'],
    [
        ('NFR-01', 'Performance', 'Page load under 3 seconds; AI recommendations under 1 second; real-time messages under 0.5 seconds.'),
        ('NFR-02', 'Security', 'JWT authentication, Row Level Security, HTTPS, identity verification, and login lockout enforced throughout.'),
        ('NFR-03', 'Availability', '99.5% uptime via Supabase cloud infrastructure on AWS with automatic failover.'),
        ('NFR-04', 'Scalability', 'Supports 10,000+ concurrent users via Vercel serverless functions and Supabase connection pooling.'),
        ('NFR-05', 'Usability', 'Mobile-responsive design tested on Chrome, Firefox, and Safari; WCAG 2.1 Level AA accessibility compliance.'),
        ('NFR-06', 'Maintainability', 'TypeScript throughout; modular Next.js App Router structure; complete API documentation.'),
        ('NFR-07', 'Data Integrity', 'PostgreSQL constraints, foreign keys, and Zod input validation at all API boundaries.'),
    ],
    caption='Table 19: Non-Functional Requirements'
)
doc.add_page_break()

# ===================== CHAPTER 3 =====================
chapter('Chapter 3')
chapter('Requirements Analysis and Design of the New System')

section('Introduction')
para('This chapter presents the detailed analysis and design of PetMatchAI. Following the Object-Oriented Methodology and the Unified Modeling Language (UML) as the primary design notation, the chapter covers use case diagrams and descriptions, class diagram, sequence diagrams, an activity diagram, the entity relationship diagram (ERD), comprehensive data dictionaries for all database tables, and the system architectural design. Together, these artefacts provide a complete blueprint of the system\'s structure, behaviour, and data architecture.')

section('Analysis and Design Methodology')
subsection('Object-Oriented Methodology')
para('PetMatchAI was designed using the Object-Oriented Methodology (OOM), which models the system as a collection of interacting objects, each encapsulating both data (attributes) and behaviour (methods). OOM was selected because it aligns naturally with the component-based architecture of Next.js, the row-based data model of PostgreSQL, and the service-oriented design of the API layer. Key OOM principles applied include encapsulation (API routes expose only necessary data), inheritance (shared UI components), and polymorphism (role-based dashboard rendering).')
subsection('Unified Modeling Language')
para('UML (Unified Modeling Language) is used as the standard notation for expressing system design across structural and behavioural dimensions. The following UML diagram types are used in this chapter: Use Case Diagrams to model system functionality from the user perspective; Class Diagrams to model the static structure of system entities and their relationships; Sequence Diagrams to model the dynamic interactions between system components for key flows; Activity Diagrams to model the workflow of the main user journey; and Entity Relationship Diagrams to model the database structure and table relationships.')

section('Use Case Diagram')
placeholder(3, 'Use Case Diagram — PetMatchAI')
para('PetMatchAI defines three primary actors: the Buyer, the Seller, and the Administrator. Each actor interacts with specific system use cases as described below:')
para('Buyer use cases include: Register Account, Login (email/password or Google OAuth), Set Preferences, Browse/Search Listings, View AI Recommendations, Save Listing, Compare Pets, Send Message, Make Offer, Leave Review, File Dispute, Manage Profile, and Configure Notifications.')
para('Seller use cases include: Register Account, Login, Create Pet Listing, Edit/Delete Listing, Respond to Messages, Accept/Reject/Counter Offer, View Listing Analytics, Submit Identity Verification, View Dashboard, and Respond to Reviews.')
para('Administrator use cases include: Manage Users (change role, suspend), Moderate Listings, Approve/Reject Identity Verifications, Resolve Disputes, View Audit Log, Broadcast Announcements, and View Platform Analytics.')

section('Use Case Descriptions')
use_case_tbl(1, 'Register User Account', 'Buyer / Seller',
    'User has a valid email address and internet access.',
    '1. User navigates to /auth/register.\n2. User selects role (Buyer or Seller).\n3. User enters name, email, phone number, and password.\n4. Password strength indicator validates the entry.\n5. For Buyer role, user completes Step 3: species preferences, budget range, and purpose.\n6. System calls supabase.auth.signUp() and creates auth.users record.\n7. Database trigger fires, inserting row in profiles table.\n8. Brevo sends OTP verification email.\n9. User confirms email at /auth/confirm-email.\n10. User redirected to role-appropriate dashboard.',
    'If email already exists, system returns "Email already registered" error. If OTP expires, user may request resend.',
    'User account created with selected role. Profile record exists in profiles table. Buyer preferences saved to buyer_preferences table (for buyer role).')

use_case_tbl(2, 'Login to System', 'Buyer / Seller / Administrator',
    'User account exists and email is verified.',
    '1. User navigates to /auth/login.\n2. User enters email and password.\n3. System calls supabase.auth.signInWithPassword().\n4. JWT session token created and stored.\n5. System checks profiles.role and redirects to appropriate dashboard.\n6. Idle timeout timer (30 minutes) started.',
    'Google OAuth available via "Sign in with Google" button (shows account picker due to prompt=select_account). After 5 failed attempts, account is locked for 15 minutes with unlock countdown displayed.',
    'Authenticated JWT session established. User redirected to role-appropriate dashboard.')

use_case_tbl(3, 'Create Pet Listing', 'Seller',
    'Seller is authenticated and has an active session.',
    '1. Seller navigates to /listings/new.\n2. Seller enters pet details: name, species, breed, age, gender, color, size, and personality traits.\n3. Seller sets health information: vaccinated, dewormed, microchipped flags; optionally uploads health certificate.\n4. Seller uploads one or more photos to Supabase Storage; optionally uploads a video.\n5. Seller enters price, location, and description.\n6. Seller optionally marks listing as Featured.\n7. On submit, POST /api/pets creates the pets record and pet_images records.\n8. Audit log entry created with action "create_listing".\n9. Matching buyers notified via /api/pets POST after handler.',
    'Seller may save draft to localStorage and resume later. Draft restored on page re-visit.',
    'Listing created with status "active" and visible on /listings. Buyers with matching preferences receive in-app and push notifications.')

use_case_tbl(4, 'Get AI Recommendations', 'Buyer',
    'Buyer is authenticated. Buyer preferences exist in buyer_preferences. Active pet listings exist.',
    '1. Buyer navigates to /recommendations.\n2. GET /api/recommendations fetches buyer_preferences and all active listings.\n3. Previously scored pets retrieved from ai_matches cache.\n4. Uncached pets sent to Groq API with buyer preferences as context; Groq returns match score (0-100) and natural language reasons array for each pet.\n5. New scores cached in ai_matches table via upsert.\n6. Pets sorted by score; top 6 returned as primary recommendations.\n7. Contextual sections computed: "Because You Viewed" (same species/breed as last viewed pet), "Similar to Your Saved Pets" (same species/breed/price range), "Popular in Your Area" (location-overlap + sorted by views).\n8. AI-generated recommendation narrative generated by Groq.\n9. Response sent to browser and rendered.',
    'If Groq API unavailable, cached scores used. If no cached scores, default score of 50 applied. If buyer has no preferences set, prompt displayed to complete profile.',
    'Personalised recommendations displayed with match percentage badges, reason tags, feedback buttons (Interested / Not Interested), and contextual sections.')

use_case_tbl(5, 'Send Message', 'Buyer / Seller',
    'Both parties are authenticated. A listing exists.',
    '1. Buyer clicks "Inquire" on listing detail page.\n2. POST /api/messages creates message_threads record (pet_id, buyer_id, seller_id) if not exists.\n3. Message inserted to messages table.\n4. Supabase Realtime subscription on message_threads fires on seller\'s connected browser.\n5. Seller sees new message in real time.\n6. Notification inserted to notifications table for offline delivery.\n7. SMS sent via Termii if seller has SMS notifications enabled.\n8. Push notification sent if seller has push enabled.\n9. Both parties can continue conversation; typing indicator broadcast via Supabase Realtime broadcast channel.',
    'If seller is offline, notification delivered on next login. Video call available via "Video Call" button which opens Jitsi Meet iframe modal.',
    'Message delivered in real time (avg 0.3 seconds). Thread updated with last message and timestamp.')

use_case_tbl(6, 'Submit Price Offer', 'Buyer',
    'Buyer is authenticated. Listing is active.',
    '1. Buyer navigates to listing detail page.\n2. Buyer clicks "Make Offer" and enters offer amount and optional note.\n3. POST /api/offers creates offers record with status "pending".\n4. Seller notified via in-app notification, email, and push.\n5. Seller navigates to /offers and views incoming offer.\n6. Seller selects Accept, Reject, or Counter.\n7. If Counter, seller enters counter amount; buyer notified.\n8. If Accepted, offer status set to "accepted"; buyer notified.\n9. Buyer can leave review after accepted offer.',
    'If offer is already pending for this pet+buyer combination, existing offer is shown. Counter-offer flow repeats until accepted or rejected.',
    'Offer status updated and persisted. Both parties notified at each status change.')

use_case_tbl(7, 'File Dispute', 'Buyer / Seller',
    'User is authenticated. An accepted offer exists.',
    '1. User navigates to /offers.\n2. User clicks "Report an Issue" on an accepted offer.\n3. Inline DisputeForm component expands with 6 preset issue type options.\n4. User selects issue type and optionally adds description.\n5. POST /api/disputes creates disputes record with status "open".\n6. All administrators notified via notifications table insert.\n7. Audit log entry created.\n8. Admin views dispute in Admin Panel > Disputes tab.\n9. Admin resolves dispute with resolution text; status updated to "resolved".\n10. Both parties notified of resolution.',
    'User may submit dispute without accepted offer (general platform dispute). Dispute status progresses: open > investigating > resolved.',
    'Dispute logged in disputes table. Admin notified. Resolution recorded when admin acts.')

section('Class Diagram')
placeholder(4, 'Class Diagram — PetMatchAI')
para('The class diagram for PetMatchAI models the principal entities and their relationships. The main classes are:')
classes = [
    ('User / Profile', 'Attributes: id (UUID), name, email, role (buyer/seller/administrator), phone, location, avatar_url, is_verified. Methods: register(), login(), logout(), updateProfile(), requestVerification().'),
    ('Pet', 'Attributes: id, seller_id, name, species, breed, age_months, gender, color, price, location, vaccinated, dewormed, microchipped, status, views, featured, image_url, traits[]. Methods: create(), edit(), delete(), incrementView(), getAnalytics().'),
    ('BuyerPreference', 'Attributes: user_id, preferred_species[], preferred_breeds[], age_min, age_max, budget_min, budget_max, preferred_location, preferred_gender, purpose, health_requirements[]. Methods: save(), update().'),
    ('AIMatch', 'Attributes: buyer_id, pet_id, score, reasons[], feedback, match_status. Methods: computeScore(), cachScore(), applyFeedback(), accept(), decline().'),
    ('MessageThread', 'Attributes: id, pet_id, buyer_id, seller_id, last_message, last_message_at, buyer_unread, seller_unread. Methods: create(), getMessages(), markRead().'),
    ('Message', 'Attributes: id, thread_id, sender_id, content, is_read, created_at. Methods: send(), markRead().'),
    ('Offer', 'Attributes: id, pet_id, buyer_id, seller_id, amount, status, counter_amount, note. Methods: submit(), accept(), reject(), counter(), notifyParties().'),
    ('Review', 'Attributes: id, reviewer_id, seller_id, pet_id, rating, communication_rating, accuracy_rating, health_rating, comment, photo_urls[], is_verified. Methods: submit(), respond(), voteHelpful().'),
    ('Notification', 'Attributes: id, user_id, type, title, message, data, is_read. Methods: send(), markRead(), markAllRead().'),
    ('Dispute', 'Attributes: id, reporter_id, respondent_id, subject, description, status, resolution, resolved_by. Methods: file(), investigate(), resolve().'),
    ('UserVerification', 'Attributes: id, user_id, id_type, id_image_url, selfie_url, status, admin_note. Methods: submit(), approve(), reject().'),
    ('AuditLog', 'Attributes: id, user_id, action, entity_type, entity_id, details. Methods: log().'),
]
for cls, desc in classes:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(cls + ': ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

section('Sequence Diagrams')
subsection('Sequence Diagram 1: User Registration Flow')
placeholder(5, 'Sequence Diagram: User Registration')
para('The user registration sequence involves the following actors and components: User Browser, Next.js Frontend (/auth/register page), Supabase Auth service, profiles database table, and Brevo Email Service.')
para('Flow: (1) User submits registration form with role, name, email, and password. (2) Frontend calls supabase.auth.signUp({email, password, options: {data: {name, role}}}). (3) Supabase Auth creates a new record in auth.users table. (4) The on_auth_user_created database trigger fires automatically. (5) Trigger inserts a new row into the profiles table with the user\'s id, email, name, and role. (6) For buyer role, frontend calls POST /api/preferences to save species preferences and budget to buyer_preferences table. (7) Brevo SMTP service sends OTP verification email to the user\'s email address. (8) User clicks verification link in email and is redirected to /auth/confirm-email. (9) Frontend calls supabase.auth.verifyOtp(). (10) On success, authenticated JWT session is established and user is redirected to role-appropriate dashboard.')
subsection('Sequence Diagram 2: AI Matching Flow')
placeholder(6, 'Sequence Diagram: AI Matching Flow')
para('The AI matching sequence involves: Buyer Browser, Next.js API Route (/api/recommendations), buyer_preferences table, pets table, Groq API, and ai_matches table.')
para('Flow: (1) Buyer browser sends GET /api/recommendations with JWT Bearer token in Authorization header. (2) API verifies JWT token via verifyToken() using Supabase admin client. (3) API fetches buyer_preferences record for the authenticated buyer. (4) API fetches buyer profile (location) from profiles table. (5) API fetches saved pets, pet view history, match feedback history, and active pet listings in parallel via Promise.all(). (6) For pets not yet scored, API constructs a structured prompt including buyer preferences, pet attributes, and feedback history, and sends batch to Groq API (llama-3.3-70b-versatile model). (7) Groq returns JSON array of {pet_id, score, reasons[]} for each unscored pet. (8) New scores are upserted to ai_matches table via after() background task. (9) All pets (cached + new) sorted by score; top 6 selected. (10) Contextual sections (becauseYouViewed, similarToSaved, popularInArea) computed from pool using lightweight heuristics. (11) Groq generates narrative recommendation message. (12) Response returned to browser with {data, message, becauseYouViewed, similarToSaved, popularInArea}.')
subsection('Sequence Diagram 3: Buyer-Seller Messaging Flow')
placeholder(7, 'Sequence Diagram: Buyer-Seller Messaging')
para('The messaging sequence involves: Buyer Browser, Next.js API Route (/api/messages), message_threads table, messages table, Supabase Realtime service, Seller Browser, notifications table, and Termii SMS service.')
para('Flow: (1) Buyer clicks "Inquire" on listing detail page; frontend calls POST /api/messages with pet_id, seller_id, and message content. (2) API verifies buyer\'s JWT token. (3) API upserts a message_threads record (pet_id, buyer_id, seller_id). (4) API inserts message into messages table with sender_id = buyer_id. (5) API updates message_threads.last_message and increments seller_unread count. (6) Supabase Realtime fires INSERT event on messages table for any subscriber watching the thread. (7) Seller\'s browser, subscribed to changes on this thread_id, receives the real-time event and renders the new message without page refresh. (8) API inserts notification record for seller (type: "message"). (9) If seller has SMS enabled in notification_prefs, Termii API called to send SMS to seller\'s Nigerian phone number. (10) If seller has push enabled, Web Push notification sent via VAPID key to seller\'s service worker.')

section('Activity Diagram')
placeholder(8, 'Activity Diagram: Main User Flow')
para('The activity diagram for PetMatchAI describes the primary buyer journey from initial platform engagement through transaction completion:')
activity_steps = [
    'START: New user arrives at PetMatchAI',
    'DECISION: Registered? → If No: Complete registration (select role, enter details, verify email)',
    'For Buyer: Complete preference setup (species, budget, purpose)',
    'LOGIN: Authenticate with email/password or Google OAuth',
    'DASHBOARD: View role-appropriate dashboard',
    'BROWSE: Navigate to /listings or /recommendations',
    'SEARCH/FILTER: Apply breed, price, location, gender filters; toggle map view',
    'VIEW LISTING: Click pet card → view listing detail with photos, health info, seller profile',
    'DECISION: Interested? → If No: Return to listings; If Yes: Choose next action',
    'SAVE: Add to saved listings for later reference; OR',
    'INQUIRE: Click "Inquire" → open messaging thread with seller',
    'COMMUNICATE: Exchange messages; negotiate; optionally initiate video call',
    'OFFER: Click "Make Offer" → enter amount → seller accepts/rejects/counters',
    'ACCEPTED OFFER: Arrange physical pet collection',
    'REVIEW: Leave multi-criteria review after transaction',
    'END: Review published; seller rating updated',
]
for step in activity_steps:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'→ {step}')
    set_font(r, size=11)

section('Database Diagram (Entity Relationship Diagram)')
placeholder(2, 'Entity Relationship Diagram (ERD) — PetMatchAI')
para('The ERD for PetMatchAI models the following entity relationships:')
erd_rels = [
    'profiles (1) — (N) pets: One seller profile may list many pets [seller_id FK]',
    'profiles (1) — (1) buyer_preferences: One buyer has one preference record [user_id FK, UNIQUE]',
    'profiles (1) — (N) saved_pets: One user may save many pets [user_id FK]',
    'pets (1) — (N) saved_pets: One pet may be saved by many users [pet_id FK]',
    'profiles (1) — (N) message_threads: One user participates in many threads [buyer_id, seller_id FK]',
    'message_threads (1) — (N) messages: One thread contains many messages [thread_id FK]',
    'pets (1) — (N) offers: One pet may receive many offers [pet_id FK]',
    'profiles (1) — (N) offers: One buyer submits many offers; one seller receives many offers [buyer_id, seller_id FK]',
    'profiles (1) — (N) reviews: One buyer writes many reviews; one seller receives many reviews [reviewer_id, seller_id FK]',
    'pets (1) — (N) reviews: One pet associated with many reviews [pet_id FK]',
    'profiles (1) — (N) ai_matches: One buyer has many AI match records [buyer_id FK]',
    'pets (1) — (N) ai_matches: One pet scored against many buyers [pet_id FK]',
    'profiles (1) — (N) notifications: One user receives many notifications [user_id FK]',
    'profiles (1) — (N) audit_logs: One user generates many audit log entries [user_id FK]',
    'profiles (1) — (N) disputes: One user files many disputes as reporter [reporter_id FK]',
    'profiles (1) — (1) user_verification_requests: One user has at most one verification request [user_id FK, UNIQUE]',
]
for rel in erd_rels:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(rel)
    set_font(r, size=11)

section('Data Dictionary')

tbl(['Column', 'Type', 'Constraint', 'Description'],
    [('id','UUID','PK','Unique identifier, references auth.users'),
     ('name','TEXT','NOT NULL','User full name'),
     ('email','TEXT','NOT NULL','User email address'),
     ('role','TEXT','CHECK (buyer/seller/administrator)','Platform role'),
     ('phone','TEXT','NULLABLE','Phone number'),
     ('location','TEXT','NULLABLE','User city/location'),
     ('avatar_url','TEXT','NULLABLE','Profile photo URL'),
     ('is_verified','BOOLEAN','DEFAULT FALSE','KYC verification status'),
     ('created_at','TIMESTAMPTZ','DEFAULT NOW()','Account creation timestamp'),
     ('updated_at','TIMESTAMPTZ','DEFAULT NOW()','Last update timestamp'),
    ],
    caption='Table 8: Data Dictionary – profiles Table')

tbl(['Column', 'Type', 'Constraint', 'Description'],
    [('id','UUID','PK','Listing unique identifier'),
     ('seller_id','UUID','FK → profiles','Seller who owns this listing'),
     ('name','TEXT','NOT NULL','Pet name'),
     ('species','TEXT','CHECK (dog/cat/other)','Pet species'),
     ('breed','TEXT','NOT NULL','Pet breed'),
     ('age_months','INTEGER','NOT NULL','Age in months'),
     ('gender','TEXT','CHECK (male/female)','Pet gender'),
     ('color','TEXT','NULLABLE','Pet colour'),
     ('price','NUMERIC','NOT NULL','Asking price in NGN'),
     ('location','TEXT','NOT NULL','Listing location'),
     ('description','TEXT','NULLABLE','Pet description'),
     ('vaccinated','BOOLEAN','DEFAULT FALSE','Vaccination status'),
     ('dewormed','BOOLEAN','DEFAULT FALSE','Deworming status'),
     ('microchipped','BOOLEAN','DEFAULT FALSE','Microchip status'),
     ('registration_info','TEXT','NULLABLE','Registration details'),
     ('pedigree','TEXT','NULLABLE','Pedigree information'),
     ('status','TEXT','CHECK (active/pending/sold)','Listing status'),
     ('views','INTEGER','DEFAULT 0','View count'),
     ('inquiries','INTEGER','DEFAULT 0','Inquiry count'),
     ('featured','BOOLEAN','DEFAULT FALSE','Featured listing flag'),
     ('image_url','TEXT','NULLABLE','Primary image URL'),
     ('traits','TEXT[]','DEFAULT {}','Personality trait tags'),
     ('size','TEXT','NULLABLE','Pet size (small/medium/large/extra_large)'),
     ('health_certificate_url','TEXT','NULLABLE','Health certificate file URL'),
    ],
    caption='Table 9: Data Dictionary – pets Table')

tbl(['Column', 'Type', 'Constraint', 'Description'],
    [('id','UUID','PK','Preference record ID'),
     ('user_id','UUID','FK → profiles, UNIQUE','Buyer owner'),
     ('preferred_species','TEXT[]','DEFAULT {}','Species preferences'),
     ('preferred_breeds','TEXT[]','DEFAULT {}','Breed preferences'),
     ('age_min','INTEGER','DEFAULT 0','Minimum age in months'),
     ('age_max','INTEGER','DEFAULT 120','Maximum age in months'),
     ('budget_min','NUMERIC','DEFAULT 0','Minimum budget NGN'),
     ('budget_max','NUMERIC','DEFAULT 999999','Maximum budget NGN'),
     ('preferred_location','TEXT','NULLABLE','Preferred location'),
     ('preferred_gender','TEXT','DEFAULT any','Gender preference'),
     ('purpose','TEXT','NULLABLE','Acquisition purpose'),
     ('health_requirements','TEXT[]','DEFAULT {}','Health requirements'),
    ],
    caption='Table 10: Data Dictionary – buyer_preferences Table')

tbl(['Column', 'Type', 'Constraint', 'Description'],
    [('id','UUID','PK','Match record ID'),
     ('buyer_id','UUID','FK → profiles','Buyer'),
     ('pet_id','UUID','FK → pets','Pet being scored'),
     ('score','INTEGER','CHECK (0–100)','AI compatibility score'),
     ('reasons','TEXT[]','DEFAULT {}','AI-generated reason tags'),
     ('feedback','TEXT','NULLABLE','User feedback: interested/not_interested'),
     ('match_status','TEXT','NULLABLE','accepted / declined'),
     ('created_at','TIMESTAMPTZ','DEFAULT NOW()','Score creation timestamp'),
     ('updated_at','TIMESTAMPTZ','DEFAULT NOW()','Last update timestamp'),
    ],
    caption='Table 11: Data Dictionary – ai_matches Table')

tbl(['Column', 'Type', 'Constraint', 'Description'],
    [('id','UUID','PK','Message ID'),
     ('thread_id','UUID','FK → message_threads','Parent thread'),
     ('sender_id','UUID','FK → profiles','Message sender'),
     ('content','TEXT','NOT NULL','Message text content'),
     ('is_read','BOOLEAN','DEFAULT FALSE','Read status'),
     ('created_at','TIMESTAMPTZ','DEFAULT NOW()','Send timestamp'),
    ],
    caption='Table 12: Data Dictionary – messages Table')

tbl(['Column', 'Type', 'Constraint', 'Description'],
    [('id','UUID','PK','Offer ID'),
     ('pet_id','UUID','FK → pets','Pet being offered on'),
     ('buyer_id','UUID','FK → profiles','Buyer making offer'),
     ('seller_id','UUID','FK → profiles','Seller receiving offer'),
     ('amount','NUMERIC','NOT NULL','Offer amount NGN'),
     ('status','TEXT','CHECK (pending/accepted/rejected/countered)','Offer status'),
     ('counter_amount','NUMERIC','NULLABLE','Counter-offer amount'),
     ('note','TEXT','NULLABLE','Optional offer note'),
     ('created_at','TIMESTAMPTZ','DEFAULT NOW()','Offer submission time'),
     ('updated_at','TIMESTAMPTZ','DEFAULT NOW()','Last status change time'),
    ],
    caption='Table 13: Data Dictionary – offers Table')

tbl(['Column', 'Type', 'Constraint', 'Description'],
    [('id','UUID','PK','Review ID'),
     ('reviewer_id','UUID','FK → profiles','Buyer leaving review'),
     ('seller_id','UUID','FK → profiles','Seller being reviewed'),
     ('pet_id','UUID','FK → pets, NULLABLE','Associated listing'),
     ('rating','INTEGER','CHECK (1–5)','Overall star rating'),
     ('communication_rating','INTEGER','CHECK (1–5), NULLABLE','Communication score'),
     ('accuracy_rating','INTEGER','CHECK (1–5), NULLABLE','Listing accuracy score'),
     ('health_rating','INTEGER','CHECK (1–5), NULLABLE','Pet health score'),
     ('comment','TEXT','NULLABLE','Written review'),
     ('is_verified','BOOLEAN','DEFAULT FALSE','Verified transaction badge'),
     ('photo_urls','TEXT[]','NULLABLE','Review photo URLs'),
    ],
    caption='Table 14: Data Dictionary – reviews Table')

tbl(['Column', 'Type', 'Constraint', 'Description'],
    [('id','UUID','PK','Notification ID'),
     ('user_id','UUID','FK → profiles','Recipient user'),
     ('type','TEXT','CHECK (match/message/offer/review/listing/price/system)','Notification type'),
     ('title','TEXT','NOT NULL','Notification title'),
     ('message','TEXT','NOT NULL','Notification body'),
     ('data','JSONB','DEFAULT {}','Additional metadata'),
     ('is_read','BOOLEAN','DEFAULT FALSE','Read status'),
     ('created_at','TIMESTAMPTZ','DEFAULT NOW()','Creation timestamp'),
    ],
    caption='Table 15: Data Dictionary – notifications Table')

tbl(['Column', 'Type', 'Constraint', 'Description'],
    [('id','UUID','PK','Dispute ID'),
     ('reporter_id','UUID','FK → profiles','User who filed dispute'),
     ('respondent_id','UUID','FK → profiles','User being disputed against'),
     ('subject','TEXT','NOT NULL','Dispute subject'),
     ('description','TEXT','NULLABLE','Detailed description'),
     ('context','TEXT','NULLABLE','Additional context'),
     ('status','TEXT','CHECK (open/investigating/resolved/closed)','Dispute status'),
     ('resolution','TEXT','NULLABLE','Admin resolution text'),
     ('resolved_by','UUID','FK → profiles, NULLABLE','Admin who resolved'),
     ('resolved_at','TIMESTAMPTZ','NULLABLE','Resolution timestamp'),
    ],
    caption='Table 16: Data Dictionary – disputes Table')

tbl(['Column', 'Type', 'Constraint', 'Description'],
    [('id','UUID','PK','Verification request ID'),
     ('user_id','UUID','FK → profiles, UNIQUE','User submitting verification'),
     ('id_type','TEXT','CHECK (national_id/passport/drivers_license/cac_certificate)','Document type'),
     ('id_image_url','TEXT','NOT NULL','Document photo URL'),
     ('selfie_url','TEXT','NULLABLE','Optional selfie URL'),
     ('status','TEXT','CHECK (pending/approved/rejected)','Review status'),
     ('admin_note','TEXT','NULLABLE','Admin rejection reason'),
     ('reviewed_at','TIMESTAMPTZ','NULLABLE','Review timestamp'),
     ('reviewed_by','UUID','FK → profiles, NULLABLE','Admin who reviewed'),
    ],
    caption='Table 17: Data Dictionary – user_verification_requests Table')

section('System Architectural Design')
placeholder(9, 'System Architecture Diagram — PetMatchAI')
para('PetMatchAI follows a three-tier client-server architecture deployed as a cloud-native application with serverless compute and managed database infrastructure.')
subsection('Presentation Layer')
para('The presentation layer is implemented using Next.js 16 with the App Router, React, and Tailwind CSS. All user-facing pages are React Server Components or Client Components with the "use client" directive where interactivity is required. The presentation layer handles routing, form rendering, state management, and real-time UI updates. Supabase client-side subscriptions provide live data updates for messaging and notifications without page refreshes. The presentation layer is deployed on Vercel\'s Edge Network, providing automatic CDN distribution and sub-100ms response times for static assets.')
subsection('Application Layer')
para('The application layer consists of Next.js API Routes, implemented as serverless functions that handle all business logic, data validation, authentication verification, and external service integration. Key API routes include /api/pets (listing CRUD), /api/messages (messaging), /api/offers (negotiation), /api/recommendations (AI scoring), /api/analytics (metrics), /api/admin (administration), /api/disputes (dispute management), and /api/verify (KYC). All API routes use Zod for input validation at system boundaries and the Supabase Admin client with the Service Role Key for database operations that bypass RLS policies when required by administrative functions.')
subsection('Data Layer')
para('The data layer consists of Supabase PostgreSQL (version 15), Supabase Storage, and Supabase Realtime. The PostgreSQL database implements Row Level Security (RLS) policies that enforce data isolation at the database level, ensuring buyers can only read their own preferences, sellers can only edit their own listings, and all sensitive data is protected even if the API layer is compromised. Supabase Storage provides S3-compatible object storage for pet images, health certificates, review photos, and KYC documents, organised into versioned buckets with signed URL access.')
subsection('External Services')
ext_services = [
    ('Groq API (llama-3.3-70b-versatile)', 'AI-powered recommendation scoring, natural language recommendation messages, and analytics executive summaries. Selected for ultra-fast inference (400+ tokens/second) enabling real-time AI responses.'),
    ('Brevo SMTP', 'Transactional email delivery for OTP verification, offer notifications, and weekly analytics reports. Configured with custom SMTP credentials in environment variables.'),
    ('Termii', 'Nigerian SMS gateway for real-time transaction alerts to buyers and sellers. Supports Nigerian phone number normalisation and Sender ID configuration.'),
    ('Web Push API / VAPID', 'Browser push notifications delivered via service worker (public/sw.js) using Voluntary Application Server Identification (VAPID) keys for authenticated push subscription.'),
    ('Jitsi Meet', 'Video calling embedded as an iframe modal in the messaging interface. Room names derived from thread IDs provide private, end-to-end encrypted video sessions with no account creation required.'),
    ('Nominatim / OpenStreetMap', 'Free geocoding API for converting pet listing location strings to coordinates for the Leaflet map view on the listings page.'),
    ('Google OAuth', 'Social sign-in implemented via Supabase Auth OAuth integration with Google Cloud Console credentials. Configured with prompt=select_account to enforce account picker on every login.'),
]
for service, desc in ext_services:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(service + ': ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)
doc.add_page_break()

# ===================== CHAPTER 4 =====================
chapter('Chapter 4')
chapter('Implementation of the New System')

section('Introduction')
para('This chapter presents the practical implementation of PetMatchAI, describing the technologies and tools used in development, providing annotated screenshots of each major interface, and documenting the software testing approach and results. The system was developed incrementally over a structured development period, with each module built, tested, and refined before proceeding to the next. The complete system was implemented as a web-based application accessible on desktop and mobile browsers, with no native mobile app required due to the responsive design and Progressive Web App capabilities of the platform.')

section('Technologies and Tools Used')
subsection('Frontend Technologies')
tech_frontend = [
    ('Next.js 16 (App Router)', 'The primary web framework, providing file-based routing through the app/ directory, React Server Components for performance, Client Components for interactivity, and API Routes as serverless functions in a single unified codebase. Next.js 16 introduces streaming, Suspense boundaries for loading states, and improved build performance over previous versions.'),
    ('TypeScript', 'Strongly typed superset of JavaScript used throughout the entire codebase. TypeScript provides compile-time type checking, superior IDE autocompletion, and improved code maintainability — particularly valuable in a large multi-module system where type safety prevents class of runtime errors.'),
    ('Tailwind CSS', 'Utility-first CSS framework enabling rapid responsive UI development without writing custom CSS files. All components use Tailwind utility classes, ensuring visual consistency and mobile-responsiveness across all screen sizes.'),
    ('Lucide React', 'Open-source icon library providing consistent, accessible SVG icons throughout the user interface. Used for navigation, action buttons, status indicators, and feature icons.'),
    ('Leaflet.js', 'Open-source JavaScript mapping library used to render the interactive map on the /listings page. Pet listing locations are geocoded via the Nominatim OpenStreetMap API and displayed as clickable markers with popup cards.'),
    ('xlsx', 'JavaScript library for generating Excel (.xlsx) workbooks. Used to export analytics data (3-sheet workbook: Summary, Listings, Market Pricing) and match history data from the matchmaking page.'),
]
for name, desc in tech_frontend:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(name + ': ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

subsection('Backend and Database Technologies')
tech_backend = [
    ('Supabase', 'Open-source Firebase alternative providing a managed PostgreSQL database, built-in authentication, real-time WebSocket subscriptions, and S3-compatible file storage. Hosted on AWS infrastructure with 99.9% uptime SLA and global CDN. Supabase was selected for its comprehensive feature set, generous free tier, Nigerian developer community support, and tight integration with Next.js.'),
    ('PostgreSQL 15', 'The relational database engine underlying Supabase. Row Level Security (RLS) policies are defined at the database level, ensuring data isolation is enforced regardless of application-layer vulnerabilities. Foreign key constraints, CHECK constraints, and unique constraints maintain data integrity throughout the 17-table schema.'),
    ('Supabase Realtime', 'WebSocket-based subscription service built into Supabase, enabling real-time data synchronisation. Used for live message delivery, typing indicators, notification badges, and review updates — all without polling or custom WebSocket server infrastructure.'),
    ('Supabase Storage', 'S3-compatible object storage for pet listing images, health certificate files, review photos, and KYC identity documents. Storage buckets are configured with Row Level Security to ensure users can only access their own files.'),
    ('Supabase Auth', 'JWT-based authentication service supporting email/password registration with email OTP verification, Google OAuth via OpenID Connect, PKCE flow for password reset, and session management with configurable timeout.'),
]
for name, desc in tech_backend:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(name + ': ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

subsection('AI and Machine Learning')
para('The Groq API with the llama-3.3-70b-versatile large language model powers three AI features in PetMatchAI:')
ai_uses = [
    'Batch Match Scoring: Pet listings are scored against buyer preferences through a structured JSON prompt that provides buyer preferences (species, breed, age range, budget, location, gender, purpose, health requirements), each pet\'s attributes, and the buyer\'s feedback history (interested/not_interested on previous matches). Groq returns a JSON array of {pet_id, score, reasons[]} objects, which are cached in the ai_matches table to minimise API calls on subsequent loads.',
    'Recommendation Narrative: A natural language recommendation message is generated summarising why the top recommended pets are a good match for the specific buyer, personalising the recommendations dashboard with an AI-written introduction.',
    'Analytics Executive Summary: Administrators can generate a 3-sentence AI executive summary of platform metrics (total listings, active users, match rates, top breeds) on the analytics dashboard, providing an instant narrative overview of platform health.',
]
for use in ai_uses:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(use)
    set_font(r, size=12)
para('Groq was selected over OpenAI and other LLM providers because of its ultra-fast inference speed (exceeding 400 tokens per second), which enables AI recommendation responses within sub-second latency on the web interface. This performance characteristic is critical for a marketplace where buyers expect immediate, responsive interactions.')

subsection('Notification Services')
notif_services = [
    ('Brevo SMTP', 'Transactional email delivery for OTP verification emails, offer status notifications, weekly analytics reports (delivered via Vercel Cron on Monday mornings), and system broadcasts. Brevo was selected for its generous free tier (300 emails/day) and reliable Nigerian email delivery rates.'),
    ('Termii', 'Nigerian SMS gateway for real-time SMS alerts on new messages and offer events. The lib/sms.ts module normalises Nigerian phone numbers (adding +234 prefix as needed) before sending. SMS delivery is gated by the user\'s notification_prefs.channel_sms preference.'),
    ('Web Push API with VAPID', 'Browser push notifications delivered via a registered service worker (public/sw.js). When users enable push notifications in Profile > Notifications, their PushSubscription object is stored in the push_subscriptions table and used to deliver push payloads via the web-push Node.js library with VAPID authentication.'),
]
for name, desc in notif_services:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(name + ': ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

subsection('Development Tools')
dev_tools = [
    ('Visual Studio Code', 'Primary IDE with TypeScript language server, ESLint, and Tailwind CSS IntelliSense extensions.'),
    ('Git and GitHub', 'Version control with feature branches and commit-based incremental development.'),
    ('Node.js 20 LTS', 'JavaScript runtime for local development and build processes.'),
    ('npm', 'Package manager for dependency management.'),
    ('Jest', 'Unit testing framework used for API route and utility function testing.'),
    ('Vercel', 'Deployment platform providing automatic CI/CD from GitHub commits, serverless function execution, and Edge Network distribution.'),
]
for name, desc in dev_tools:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(name + ': ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

section('Screenshots of the Implemented System')
screenshots = [
    (10, 'Screenshot: Login Page', 'The login page at /auth/login presents email/password fields, a "Sign in with Google" button integrated with Google OAuth (configured via Supabase Auth), a "Forgot Password" link initiating the PKCE-based reset flow, and a link to the registration page. The page enforces login lockout after five consecutive failed attempts, displaying the remaining lockout time.'),
    (11, 'Screenshot: Buyer Dashboard', 'The buyer dashboard at /dashboard/buyer displays personalised AI recommendation preview cards fetched from /api/recommendations, saved listings, recent search chips from localStorage, and follow-up reminders for accepted matches with no recent message activity. Quick action buttons navigate to /listings, /recommendations, and /messages.'),
    (12, 'Screenshot: Pet Listings Page', 'The /listings page displays a responsive grid of PetCard components. The search bar with breed autocomplete (activating at 2+ characters), species type filters, gender pills (Male/Female/Any), price range input, sort options (Newest/Price/Age/Most Popular), and a "Verified Sellers Only" toggle are shown in the filter panel. A Map toggle renders an interactive Leaflet map with geocoded listing markers.'),
    (13, 'Screenshot: AI Recommendations Page', 'The /recommendations page shows the top 6 AI-scored pets with green/amber/red match percentage badges, natural language recommendation message from Groq, "Why this pet" reason tags, Interested/Not Interested feedback buttons, Accept/Decline action buttons, and three contextual sections (Because You Viewed, Similar to Your Saved Pets, Popular in Your Area). The History tab shows all previously scored pets.'),
    (14, 'Screenshot: Matchmaking Page', 'The /matchmaking page displays pending AI matches as cards showing the pet photo, name, breed, price, and a 5-cell compatibility grid (Species, Price, Breed, Location, Health) with green checkmarks or red crosses for each dimension. Accept and Decline buttons update match_status. The History tab shows all accepted and declined matches with score and status badges.'),
    (15, 'Screenshot: Messages Page', 'The /messages page shows the thread sidebar with search, conversation thread view with timestamped messages, read receipt indicators (double checkmark), animated typing indicator bubble when the other party is typing, a Video Call button opening the Jitsi Meet iframe modal, quick reply template selector (5 pre-written messages), and archive/restore controls on thread hover.'),
    (16, 'Screenshot: Analytics Dashboard', 'The /analytics page shows KPI summary cards (Total Listings, Active Users, Average Price, Conversion Rate, Average Days to Sell, Total Offers, Total Reviews), a Breed Price Distribution bar chart, a Listings Volume by Month bar chart, a Seller Performance section, an AI Executive Summary card with Generate/Regenerate button, and a Custom Report Builder with date range, species, and status filters plus Excel export.'),
    (17, 'Screenshot: Admin Panel', 'The /admin page shows an 8-tab administration interface: Users (table with role dropdown — buyer/seller only — and Suspend button), Listings (all listings with status management), Verifications (KYC document queue showing ID type badge, document thumbnail, Approve/Reject buttons with admin note), Reviews (moderation queue), Audit Log (filterable by date range/action/user ID with CSV export), Disputes (table with inline resolution form), Categories (species chart and searchable breed directory with links to filtered listings), and AI Cache management.'),
    (18, 'Screenshot: Profile and Identity Verification', 'The /profile page shows tabs for Profile (edit name, phone, location, avatar), Preferences (buyer species, breed, budget settings), Notifications (channel toggles for email/SMS/push, event type preferences), and Reviews (submitted review history). The identity verification flow at /profile/verify shows a 3-step process: Step 1 selects ID type (National ID, Passport, Driver\'s License, CAC Certificate); Step 2 uploads document photo; Step 3 optionally uploads selfie. Pending, approved, and rejected states each display appropriate status UI.'),
    (19, 'Screenshot: Offers and Dispute Filing', 'The /offers page shows tabs for incoming and outgoing offers with status badges (Pending/Accepted/Rejected/Countered), offer amounts, and action buttons (Accept, Reject, Counter for sellers; Leave Review and Report an Issue for buyers on accepted offers). The Report an Issue button expands an inline DisputeForm with 6 preset issue categories and an optional description field.'),
]
for fig_num, label, desc in screenshots:
    subsection(label)
    para(desc)
    placeholder(fig_num, label)
    doc.add_paragraph()

section('Software Testing')
para('PetMatchAI was tested through a combination of incremental unit testing, integration testing against the live Supabase development environment, and manual user acceptance testing (UAT) with test buyer, seller, and administrator accounts. Each module was tested upon completion before development proceeded to the next increment.')
subsection('Unit Testing')
para('Unit tests were written using Jest for utility functions in the lib/ directory, including the match scoring logic in lib/groq.ts, phone number normalisation in lib/sms.ts, and input validation schemas in the API route files. Tests verified that boundary conditions — empty preference arrays, zero budgets, missing location data — were handled gracefully without exceptions.')
subsection('Integration Testing')
para('Integration tests were conducted against the live Supabase test environment (a separate project with identical schema). API routes were tested using direct HTTP requests with test JWT tokens, verifying that database operations, RLS policies, and notification triggers functioned correctly end-to-end. The real-time subscription behaviour was tested by running simultaneous buyer and seller browser sessions.')
subsection('User Acceptance Testing')
para('Manual UAT was conducted with three test user accounts: one buyer (with preferences for dogs, Lagos location, NGN 100,000 budget), one seller (with three active listings), and one administrator. The complete buyer journey — registration through review — was exercised for each core scenario listed in Table 21.')

tbl(['#', 'Test Case', 'Module', 'Input', 'Expected Output', 'Actual Output', 'Status'],
    [
        ('1','User Registration','Auth','Valid email, password, buyer role','Account created, OTP email sent, preferences captured','Account created; OTP received within 30 seconds; preferences saved','PASS'),
        ('2','Google OAuth Login','Auth','Click Google Sign-In button','Google account picker shown; redirect to dashboard','Account picker shown; redirected to buyer dashboard after consent','PASS'),
        ('3','Login Lockout','Auth','5 consecutive wrong passwords','Account locked for 15 minutes; unlock countdown shown','Lockout activated at 5th attempt; countdown displayed','PASS'),
        ('4','Create Pet Listing','Listing','Valid pet data + 3 photos','Listing created active; matching buyers notified','Listing visible on /listings; 2 test buyers received push notification','PASS'),
        ('5','AI Recommendation','AI Engine','Buyer: dog preference, NGN 80K budget','Scored recommendations; target breed scores > 70%','Average 78% match score for preferred breed listings','PASS'),
        ('6','Real-time Messaging','Communication','Buyer sends message in thread','Message delivered < 1 second; seller sees it without refresh','Delivered in 0.3 seconds via Supabase Realtime','PASS'),
        ('7','Make Offer','Offers','Offer amount below asking price','Offer submitted; seller notified via push','Seller received push notification within 2 seconds','PASS'),
        ('8','Submit Review','Reviews','4-star rating + comment after accepted offer','Review saved; aggregate rating updated on listing','Review displayed; listing rating updated to reflect new average','PASS'),
        ('9','File Dispute','Disputes','Select issue type + description','Dispute logged; admin notified','Admin received in-app notification; dispute visible in admin panel','PASS'),
        ('10','Identity Verification','KYC','Upload government ID photo (JPEG)','Request submitted; status shows Under Review','Status changed to pending; admin verification queue shows request','PASS'),
        ('11','Admin Audit Log Filter','Admin','Filter: action = create_listing','Filtered audit log rows returned','Correct entries returned; CSV export generated correctly','PASS'),
        ('12','Export Analytics','Analytics','Click Export Excel button','3-sheet Excel workbook downloaded','Excel file generated with Summary, Listings, and Market Pricing sheets','PASS'),
    ],
    caption='Table 21: Software Testing Results'
)

subsection('Hardware and Software Requirements')
tbl(['Category', 'Component', 'Specification'],
    [('Client-Side','Web Browser','Chrome 90+, Firefox 88+, Safari 14+, or Edge 90+'),
     ('Client-Side','RAM','Minimum 4 GB'),
     ('Client-Side','Internet Connection','Minimum 1 Mbps broadband'),
     ('Client-Side','Display','Minimum 320px width (mobile-responsive)'),
     ('Server-Side','Database','Supabase PostgreSQL 15 (managed AWS infrastructure)'),
     ('Server-Side','Compute','Vercel Serverless (Node.js 20, auto-scaling)'),
     ('Server-Side','Storage','Supabase Storage (S3-compatible, managed)'),
     ('Server-Side','RAM (serverless)','Auto-allocated by Vercel (up to 1 GB per function)'),
     ('Server-Side','Bandwidth','Unlimited via Vercel Edge Network CDN'),
    ],
    caption='Table 20: Hardware and Software Requirements'
)
doc.add_page_break()

# ===================== CHAPTER 5 =====================
chapter('Chapter 5')
chapter('Conclusion and Recommendations')

section('Conclusion')
para('This project has successfully achieved its primary objective: the design and implementation of an AI-powered buyer-seller matchmaking system that intelligently connects pet buyers and sellers based on personalised preferences and behavioural data at Hemmy Kennel, Lagos, Nigeria. PetMatchAI replaces a fragmented, WhatsApp-based manual process with a centralised, intelligent, and scalable digital platform comprising thirteen fully integrated functional modules.')
para('Each of the eight specific objectives has been met. A structured relational database of seventeen tables — implemented in Supabase PostgreSQL with Row Level Security — provides the data foundation for intelligent matching, search, communication, and analytics. The AI recommendation engine, powered by the Groq LLM API (llama-3.3-70b-versatile), delivers personalised match scores with transparent percentage indicators and natural language explanations, demonstrating that large language models can be effectively applied to structured buyer-seller matching in specialised domains beyond the well-studied cases of film and retail recommendation.')
para('The intelligent search and filtering module — featuring breed autocomplete, species and gender filters, price range controls, map-based location discovery via Leaflet and OpenStreetMap, and a Verified Sellers Only toggle — provides buyers with precise pet discovery capabilities that were entirely absent from the previous manual process. The real-time communication system delivers messages in an average of 0.3 seconds via Supabase Realtime WebSocket subscriptions, with typing indicators, read receipts, message templates, and Jitsi Meet video calling providing a complete buyer-seller interaction suite.')
para('The analytics and decision support module delivers KPI dashboards with listing conversion rates, average time-to-sell metrics, breed popularity rankings, geographic demand data, and AI-generated executive summaries — addressing the business intelligence gap that prevented Hemmy Kennel from making data-driven decisions about breeding priorities, pricing strategy, and inventory management. The multi-channel notification system, spanning in-app notifications, Brevo email delivery, Termii SMS for Nigerian phone numbers, and Web Push via VAPID keys, ensures that buyers and sellers remain engaged with platform activity regardless of their current device or connectivity state.')
para('The security and audit architecture — combining JWT authentication, Row Level Security at the database level, identity verification via government ID upload, login lockout after five failed attempts, 30-minute idle session timeout, and comprehensive audit logging — provides a robust and accountable platform that meets the trust and privacy expectations of both buyers and sellers.')
para('Testing across all thirteen modules confirmed that the system meets all functional and non-functional requirements. All twelve documented test cases passed, with AI scoring, real-time messaging, push notifications, Excel export, audit log filtering, and KYC verification all functioning as specified. The implemented system is ready for production deployment and represents a complete digital transformation of Hemmy Kennel\'s buyer-seller engagement model.')
para('Beyond its immediate practical impact at Hemmy Kennel, this project contributes to the academic literature by demonstrating a practical architecture for LLM-powered buyer-seller matchmaking in a niche domain, integrating AI inference, real-time communication, multi-channel notification, and comprehensive audit capability within a single serverless web application stack. The combination of Next.js 16, Supabase, and Groq proved effective for rapid, cost-efficient development of an intelligent marketplace platform, providing a replicable pattern for similar domain-specific matchmaking systems.')

section('Recommendations')
recs = [
    ('Payment Gateway Integration', 'Integrating a payment gateway such as Flutterwave or Paystack would enable secure in-platform payment processing, reducing the risk associated with offline cash transactions and providing a complete, end-to-end digital transaction record. This is the highest-priority enhancement for commercial deployment.'),
    ('React Native Mobile Application', 'Developing a React Native mobile app using the existing PetMatchAI API layer would extend platform reach to mobile-native users, enabling push notifications at the OS level, camera integration for listing photo capture, and GPS-based location search without browser permissions.'),
    ('Custom Machine Learning Model', 'As the ai_matches table accumulates match outcome data (accepted/declined feedback, completed transactions, review ratings), a custom collaborative filtering or neural recommendation model can be trained on this proprietary dataset, potentially improving match accuracy beyond the LLM-based approach used in the current implementation.'),
    ('Expanded Breeder Network', 'Extending PetMatchAI from a single-business platform to a multi-seller marketplace would dramatically increase listing volume and buyer utility. This would require seller onboarding workflows, subscription tiers, and marketplace trust and safety policies.'),
    ('Demand Forecasting Module', 'Seasonal breed demand prediction using historical listing view, inquiry, and sale data would enable Hemmy Kennel to anticipate demand cycles, adjust breeding programmes, and optimise inventory. This feature requires sufficient historical data accumulation over at least 12 months of platform operation.'),
    ('Video Listing Feature', 'Allowing sellers to upload short pet behaviour videos as part of their listing profile — beyond the current static photo gallery — would significantly improve buyer confidence and reduce the need for physical visits before purchase decisions.'),
]
for title, desc in recs:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(title + ': ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)
doc.add_page_break()

# ===================== REFERENCES =====================
chapter('References')
para('Books', bold=True)
book_refs = [
    'Aggarwal, C. C. (2016). Recommender Systems: The Textbook. Springer International Publishing.',
    'Sommerville, I. (2016). Software Engineering (10th ed.). Pearson Education.',
    'Pressman, R. S., & Maxim, B. R. (2019). Software Engineering: A Practitioner\'s Approach (9th ed.). McGraw-Hill Education.',
    'Tanenbaum, A. S., & Van Steen, M. (2017). Distributed Systems: Principles and Paradigms (3rd ed.). CreateSpace Independent Publishing.',
]
for ref in book_refs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    r = p.add_run(ref)
    set_font(r, size=12)

para('Journals and Conference Papers', bold=True)
journal_refs = [
    'Burke, R. (2002). Hybrid recommender systems: Survey and experiments. User Modeling and User-Adapted Interaction, 12(4), 331-370.',
    'Chen, L., & Pu, P. (2012). Critiquing-based recommenders: Survey and emerging trends. User Modeling and User-Adapted Interaction, 22(1-2), 125-150.',
    'Covington, P., Adams, J., & Sargin, E. (2016). Deep neural networks for YouTube recommendations. In Proceedings of the 10th ACM Conference on Recommender Systems (pp. 191-198). ACM.',
    'Grewal, D., Hulland, J., Kopalle, P. K., & Karahanna, E. (2020). The future of technology and marketing: A multidisciplinary perspective. Journal of the Academy of Marketing Science, 48(1), 1-8.',
    'Jørgensen, M., & Shepperd, M. (2007). A systematic review of software development cost estimation studies. IEEE Transactions on Software Engineering, 33(1), 33-53.',
    'Linden, G., Smith, B., & York, J. (2003). Amazon.com recommendations: Item-to-item collaborative filtering. IEEE Internet Computing, 7(1), 76-80.',
    'Resnick, P., & Varian, H. R. (1997). Recommender systems. Communications of the ACM, 40(3), 56-58.',
    'Touvron, H., Lavril, T., Izacard, G., Martinet, X., Lachaux, M., Lacroix, T., ... & Lample, G. (2023). LLaMA: Open and efficient foundation language models. arXiv preprint arXiv:2302.13971.',
    'Verma, S., & Sharma, R. (2020). Artificial intelligence in marketing: Systematic review and future research directions. International Journal of Information Management Data Insights, 1(1), 100002.',
    'Zhang, S., Yao, L., Sun, A., & Tay, Y. (2019). Deep learning based recommender system: A survey and new perspectives. ACM Computing Surveys, 52(1), 1-38.',
]
for ref in journal_refs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    r = p.add_run(ref)
    set_font(r, size=12)

para('Online Sources and Technical Documentation', bold=True)
web_refs = [
    'Groq Inc. (2024). Groq API Documentation. Retrieved from https://console.groq.com/docs',
    'Next.js. (2024). Next.js Documentation: App Router. Vercel Inc. Retrieved from https://nextjs.org/docs',
    'OWASP Foundation. (2021). OWASP Top 10: The Ten Most Critical Web Application Security Risks. Retrieved from https://owasp.org/Top10/',
    'PostgreSQL Global Development Group. (2024). PostgreSQL 15 Documentation. Retrieved from https://www.postgresql.org/docs/',
    'Supabase Inc. (2024). Supabase Documentation: Authentication, Database, and Realtime. Retrieved from https://supabase.com/docs',
    'Vercel Inc. (2024). Vercel Platform Documentation. Retrieved from https://vercel.com/docs',
    'Nigeria Pet Industry Overview. (2023). Pet Care Industry in Nigeria: Market Analysis Report. Lagos Business School Digital Economy Report.',
]
for ref in web_refs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    r = p.add_run(ref)
    set_font(r, size=12)
doc.add_page_break()

# ===================== APPENDICES =====================
chapter('Appendices')
section('Appendix A: Curriculum Vitae')
cv_data = [
    ('Full Name', 'FAMILONI Emmanuel Eniola'),
    ('Date of Birth', '[To be completed]'),
    ('Nationality', 'Nigerian'),
    ('Email Address', 'Famhemmy3@gmail.com'),
    ('Phone Number', '+250792525543'),
    ('Current Institution', 'Adventist University of Central Africa (AUCA)'),
    ('Programme', 'Bachelor of Science in Information Technology — Software Engineering'),
    ('Expected Graduation', 'June 2026'),
]
tbl(['Field', 'Details'], cv_data)

para('Technical Skills', bold=True)
skills = ['JavaScript / TypeScript', 'React / Next.js', 'Node.js', 'PostgreSQL / Supabase', 'REST API Development', 'AI/ML Integration (Groq, LLM APIs)', 'Tailwind CSS', 'Git / GitHub', 'Vercel Deployment', 'Mobile-Responsive Web Design']
for skill in skills:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(skill)
    set_font(r, size=12)

doc.add_paragraph()
section('Appendix B: Data Collection Authorization Letter')
placeholder_p = doc.add_paragraph()
shd2 = OxmlElement('w:shd')
shd2.set(qn('w:val'), 'clear')
shd2.set(qn('w:color'), 'auto')
shd2.set(qn('w:fill'), 'EEEEEE')
placeholder_p._p.get_or_add_pPr().append(shd2)
placeholder_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = placeholder_p.add_run('[Data Collection Authorization Letter — attach signed letter from Hemmy Kennel management here]')
set_font(r2, size=11, italic=True, color=(100,100,100))

doc.save(r'C:/Users/PC/petmatchai/Familoni_Thesis_25951.docx')
print("SUCCESS: Thesis saved to C:/Users/PC/petmatchai/Familoni_Thesis_25951.docx")
